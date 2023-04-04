from docx import Document, shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class Word:
    def __init__(self, path: str = None, text: str = None) -> None:
        self.__path = path
        self.__text = text

    def create_document(self) -> None:
        self.doc = Document()

    def add_style_at_document(self) -> None:
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = shared.Pt(14)

    def add_paragraph_in_document(self) -> None:
        paragraph = self.doc.add_paragraph(self.__text)
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def save_document(self) -> None:
        try:
            self.doc.save(f"{self.__path}/Результат.docx")
        except PermissionError:
            pass
