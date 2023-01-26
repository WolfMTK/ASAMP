from docx import Document, shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocumentWord:
    doc: Document

    def __init__(self, path: str = None, string: str = None):
        self.path = path
        self.string = string

    def run_script(self):
        """Запустить скрипт."""
        self.create_document()
        self.add_style_at_document()
        self.add_paragraph_in_document()
        self.save_document()

    def create_document(self):
        """Создать документ"""
        self.doc = Document()

    def add_style_at_document(self):
        """Добавить стиль к документу."""
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = shared.Pt(14)

    def add_paragraph_in_document(self):
        """Добавить параграф в документ."""
        paragraph = self.doc.add_paragraph(self.string)
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def save_document(self):
        """Сохранить документ."""
        try:
            self.doc.save(f"{self.path}/Результат.docx")
        except PermissionError:
            pass


if __name__ == "__main__":
    DocumentWord("./", "213214 text").run_script()
